from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_ORIENTATION
from docx.shared import Pt,Inches,Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
import pandas as pd

def uploadTimeTable(table,l1,l2):
    r1c1=table.cell(0,0)
    r1c1.text="Date"
    table.rows[0].cells[0].paragraphs[0].runs[0].font.bold=True
    r1c2=table.cell(0,1)
    r1c2.text="Subject"
    table.rows[0].cells[1].paragraphs[0].runs[0].font.bold=True
    for i in range(1,7):
        rc=table.cell(i,0)
        rc.text=l1[i-1]
    for j in range(1,7):
        rc=table.cell(j,1)
        rc.text=l2[j-1]
    for i in range(7):
        for j in range(2):
            table.cell(i,j).paragraphs[0].paragraph_format.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER

doc=Document()
section=doc.sections[-1]
section.page_height=Cm(29.7)
section.page_width=Cm(21)
section.left_margin=Cm(1.27)
section.right_margin=Cm(1.27)
section.top_margin=Cm(1.27)
section.bottom_margin=Cm(0.3)

new_width,new_height=section.page_height,section.page_width
section.orientation=WD_ORIENTATION.LANDSCAPE
section.page_width=new_width
section.page_height=new_height
sectPr=section._sectPr
cols=sectPr.xpath('./w:cols')[0]
cols.set(qn('w:num'),'2')

exam=input("\nEnter the name of exam :\n")
'''d1=input("\nEnter the dates of class Nursery to UKG separated by a comma. (Ex - 12/12/2021,13/12/2021) :\n").split(',')
s1=input("\nEnter the subjects of class Nursery to UKG separated by a comma. (Ex - Maths,Hindi,English) :\n").split(',')
d2=input("\nEnter the dates of class 1st to 5th separated by a comma. :\n").split(',')
s2=input("\nEnter the subjects of class 1st to 5th separated by a comma. :\n").split(',')
d3=input("\nEnter the dates of class 6th to 8th separated by a comma. :\n").split(',')
s3=input("\nEnter the subjects of class 6th to 8th separated by a comma. :\n").split(',')'''

d1=['19/03/2022','21/03/2022','23/03/2022','24/03/2022','25/03/2022','26/03/2022']                      # dates of class Nursery to UKG
s1=['English[W]','English[O]','Hindi[W]','Hindi[O]','Maths[W]','Maths[O]']                              # subjects of class Nursery to UKG
d2=['19/03/2022','21/03/2022','23/03/2022','24/03/2022','25/03/2022','26/03/2022']                      # dates of class 1st to 5th
s2=['Ev.S.','Mathematics','English','Hindi','GK+computer','Conversation']                             # subjects of class 1st to 5th
d3=['19/03/2022','21/03/2022','23/03/2022','24/03/2022','25/03/2022','26/03/2022']                      # dates of class 6th to 8th
s3=['So. Science','Science','Mathematics','English','Hindi','Sanskrit']                              # subjects of class 6th to 8th
t1='08:30 AM - 11:00 AM'                                                                               # time of shift 1
t2='11:00 AM - 01:00 PM'                                                                               # time of shift 2

n=0
for k in range(-2,9):
    if(k==-2):
        clas='Nursery'
    elif(k==-1):
        clas='LKG'
    elif(k==0):
        clas='UKG'
    elif(k==1):
        clas='1st'
    elif(k==2):
        clas='2nd'
    elif(k==3):
        clas='3rd'
    else:
        clas=f'{k}th'
    
    if(k>=-2 and k<=0):
        l1,l2=d1,s1
        timee=t2
    elif(k>=1 and k<=5):
        l1,l2=d2,s2
        if(k!=5):
            timee=t2
        else:
            timee=t1
    else:
        l1,l2=d3,s3
        timee=t1
    
    df=pd.read_excel(f"{clas}.xlsx")
    try:
        i=0
        while(True):
            Name=df.iloc[i].Name
            Roll=df.iloc[i].RollNo
            table1=doc.add_table(rows=1,cols=1,style='Table Grid')
            table2=doc.add_table(rows=1,cols=2,style='Table Grid')
            row1=table1.rows[0]
            heading=row1.cells[0].paragraphs[0].add_run("Millennium Model School Mandi Bamora\n")
            heading.bold=True
            heading.font.size=Pt(15)
            heading1=row1.cells[0].paragraphs[0].add_run(f'Admit Card - {exam}')
            heading1.italic=True
            table1.cell(0,0).paragraphs[0].paragraph_format.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
            content=table2.cell(0,0)
            p=content.add_paragraph(f'Name      :  {Name}\nClass       :  {clas}\nRoll No.  :  {Roll}\n\n\n\t\t\t              Headmaster')
            p.paragraph_format.line_spacing=Pt(18)
            q=content.add_paragraph('\n\n')
            q.paragraph_format.line_spacing=Pt(3.2) 
            timeTable=table2.cell(0,1)
            gs=table2.cell(0,1).add_paragraph(f'Time : {timee}')
            gs.paragraph_format.space_after=Pt(3)
            gs.paragraph_format.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
            cellTable=timeTable.add_table(rows=7,cols=2)
            cellTable.alignment=WD_TABLE_ALIGNMENT.CENTER
            cellTable.style='Table Grid'
            uploadTimeTable(cellTable,l1,l2)
            #table2.cell(0,1).width=Cm(4)
            for cell in table2.columns[0].cells:
                cell.width=Inches(12)
            for cell in cellTable.columns[0].cells:
                cell.width=Inches(0.3)
            for cell in cellTable.columns[1].cells:
                cell.width=Inches(0.3)
            for cell in table2.columns[1].cells:
                cell.width=Inches(1)
            



            i=i+1
            n=n+1
            if(n%6!=0):
                doc.add_paragraph(' ')
            if(n%6==0):
                doc.add_page_break()
    except:
        pass

doc.save(f'Admit_Card_{exam}.docx')
print(f"\nYour document has been prepared with the file name : Admit_Card_{exam}.docx\n")